#!/usr/bin/env python
import argparse
import os
import sys
import time

from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from PIL import ExifTags, Image


def test_file():
    arguments = ["-d", os.getcwd() + "/new/"]
    # file_list(directory)
    pic_into_docx(arguments)


def get_args(argv):
    parser = argparse.ArgumentParser(usage="Creates a docx file containing all "
                                           "images in the specified directory")
    parser.add_argument("-d", "--directory", dest="directory",
                        help="specify directory containing image files.")
    return parser.parse_args(argv)


def date_from_system(file_path):
    file_stats = os.stat(file_path)[8]
    file_time_stamp_object = time.localtime(file_stats)
    return file_time_stamp_object


def datetime_img_list(directory_path):
    extensions = {'jpg', 'bmp', 'png', 'gif'}
    datetime_image_names = []
    for fn in os.listdir(directory_path):
        if any(fn.endswith(ext) for ext in extensions):
            datetime_image_names.append(
                (date_from_system(directory_path + fn), directory_path + fn)
            )
    return datetime_image_names


def create_date_section(run, text):
    run.add_break(WD_BREAK.PAGE)
    run.add_text(text)
    run.add_break()


def insert_picture(run, path='', set_width=3.75, line_break=True):
    # height is automatically calculated
    run.add_picture(path, width=Inches(set_width))
    if line_break:
        run.add_break()


def date_meta(image_path):
    image = Image.open(image_path)
    # exif = {ExifTags.TAGS[k]: v for k, v in image._getexif().items() if k in ExifTags.TAGS}
    # get datetime stamp, format is str '%Y:%m:%d %H:%M:%S'
    datetime_orig = ''
    for key, str_value in image._getexif().items():
        if ExifTags.TAGS[key] == 'DateTimeOriginal':
            datetime_orig = str_value
    assert datetime_orig, "{} no datetime stamp\n".format(image_path)
    time_obj = datetime.strptime(datetime_orig, "%Y:%m:%d %H:%M:%S")
    return time_obj.strftime('%B'), \
           time_obj.strftime('%d'), \
           time_obj.strftime('%Y')


def pic_into_docx(inputs=sys.argv[1:]):
    args = get_args(inputs)
    assert os.path.exists(args.directory), \
        "{} is not a valid directory".format(args.directory)

    all_images = datetime_img_list(args.directory)
    all_images.sort()

    document = Document()
    start_run = document.add_paragraph().add_run()
    date_section_orig = 0
    for file_info in all_images:
        time_obj, image_file_path = file_info
        date_section_new = time.strftime('%B %d %Y', time_obj)
        if not date_section_orig == date_section_new:
            create_date_section(start_run, date_section_new)
            date_section_orig = date_section_new
        insert_picture(start_run, image_file_path)

    file_time = datetime.now().strftime('%H%M%S')
    document.save('uscis_pictures_{}.docx'.format(file_time))


if __name__ == '__main__':
    pic_into_docx()
